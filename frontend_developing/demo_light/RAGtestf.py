# env requirment: python 3.10.x, langchain-core==0.3.37 langchain-community==0.2.4 --no-cache-dir
# mineru install: pip install -U "magic-pdf[full]" --extra-index-url https://wheels.myhloli.com -i https://mirrors.aliyun.com/pypi/simple
# mineru model download: pip install huggingface_hub
# wget https://gcore.jsdelivr.net/gh/opendatalab/MinerU@master/scripts/download_models_hf.py -O download_models_hf.py
# python download_models_hf.py
# model config adjustment: "formula-config": false, "table-config": false

import os
import re
import json
import fitz
import inspect
from dotenv import load_dotenv
import langchain_core
from typing import Any
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda
from langchain_core.output_parsers import StrOutputParser  
from typing import List, Dict, Any
from docx.shared import Inches
from magic_pdf.data.data_reader_writer import FileBasedDataWriter, FileBasedDataReader
from magic_pdf.data.dataset import PymuDocDataset
from magic_pdf.model.doc_analyze_by_custom_model import doc_analyze
from magic_pdf.config.enums import SupportedPdfParseMethod


# 保留您原有的图片处理代码
def delete_all_files_in_folder(folder_path):
    """清空目录"""
    if os.path.exists(folder_path) and os.path.isdir(folder_path):
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.isfile(file_path):
                os.remove(file_path)

def rename_images_from_json(image_folder, json_file):
    """保留您原有的重命名逻辑"""
    with open(json_file, 'r', encoding='utf-8') as file:
        data = json.load(file)
    
    figure_pattern = re.compile(r"Figure (\d+)")
    
    for index, entry in enumerate(data):
        if entry['type'] == 'image':
            img_path = entry['img_path']
            img_caption = entry['img_caption']
            
            figure_match = None
            for caption in img_caption:
                figure_match = figure_pattern.search(caption)
                if figure_match:
                    break
            
            if figure_match:
                figure_number = figure_match.group(1)
                new_filename = f"figure{figure_number}.jpg"
            else:
                new_filename = f"figure{index + 1}.jpg"
            
            current_filename = img_path.split('/')[-1]
            old_filepath = os.path.join(image_folder, current_filename)
            new_filepath = os.path.join(image_folder, new_filename)
            
            if os.path.exists(old_filepath):
                os.rename(old_filepath, new_filepath)

class EnhancedPaperAnalyzer:
    def __init__(self, pdf_path: str):
        self.pdf_path = pdf_path
        self.vectorstore = None
        self.rag_chain = None
        self.figures = {}
        self.embeddings = OpenAIEmbeddings(model="text-embedding-3-large")
        self.llm = ChatOpenAI(model="o1-mini")
        self.processed_chunks = []  # 新增实例变量存储分块列表


    def process_document(self, image_output_dir: str):
        """完整文档处理流程"""
        self._extract_images_and_text(image_output_dir)
        self._process_text_with_langchain()

    def _extract_images_and_text(self, output_dir: str):
        """保留您的magic_pdf图片提取逻辑"""
        delete_all_files_in_folder(output_dir)
        
        # 获取不含路径和后缀的文件名
        name_without_suff = os.path.splitext(os.path.basename(self.pdf_path))[0]  # 修正文件名获取方式
        
        # 生成JSON文件到上级目录
        json_path = os.path.join(os.path.dirname(output_dir), f"{name_without_suff}_content_list.json")  # 修正路径
        
        reader = FileBasedDataReader("")
        pdf_bytes = reader.read(self.pdf_path)
        ds = PymuDocDataset(pdf_bytes)
        
        image_writer = FileBasedDataWriter(output_dir)
        md_writer = FileBasedDataWriter(os.path.dirname(output_dir))  # 将JSON保存到输出目录的上级

        if ds.classify() == SupportedPdfParseMethod.OCR:
            infer_result = ds.apply(doc_analyze, ocr=True)
            pipe_result = infer_result.pipe_ocr_mode(image_writer)
        else:
            infer_result = ds.apply(doc_analyze, ocr=False)
            pipe_result = infer_result.pipe_txt_mode(image_writer)
        
        # 将content_list.json保存到指定位置
        pipe_result.dump_content_list(md_writer, f"{name_without_suff}_content_list.json", output_dir)
        
        # 使用绝对路径执行重命名
        rename_images_from_json(output_dir, json_path)  # 使用修正后的路径
        
        # 构建图片索引
        for file in os.listdir(output_dir):
            if file.startswith("figure") and file.endswith(".jpg"):
                figure_num = re.findall(r"figure(\d+)", file)
                if figure_num:
                    self.figures[f"图{figure_num[0]}"] = os.path.abspath(os.path.join(output_dir, file))

    def _process_text_with_langchain(self):
        """动态适配不同版本Document类的处理流程"""
        # 获取Document构造函数参数列表
        init_args = inspect.getfullargspec(Document.__init__).args
        
        # 根据参数名选择内容字段
        content_param = "text" if "text" in init_args else "page_content"
        
        loader = PyPDFLoader(self.pdf_path)
        pages = loader.load()
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", "(?<=\. )", " "]
        )
        
        chunks = text_splitter.split_documents(pages)
        
        processed_chunks = []
        for idx, chunk in enumerate(chunks):
            new_metadata = chunk.metadata.copy()
            new_metadata["chunk_index"] = str(idx)
            processed_chunks.append(
                Document(page_content=chunk.page_content, metadata=new_metadata)
            )
        print(f"Processed chunks metadata: {[doc.metadata for doc in processed_chunks]}")
        self.vectorstore = FAISS.from_documents(
        processed_chunks,
        self.embeddings
    )
        # 添加验证
        print(f"Vectorstore created with {len(processed_chunks)} documents")
        self.processed_chunks = processed_chunks


    def generate_analysis(self):
        """生成双索引分析报告"""
        prompt_template = """请你扮演一位专业的 AI 领域学术论文分析员，根据我提供的论文内容，
        并结合你自身的知识库背景,基于以下上下文生成分析报告：
        上下文：
        {context}

        [写作目标]:请你扮演一位专业的 AI 领域学术论文分析员，根据我提供的论文内容，并结合你自身的知识库背景，完成一份包含以下三节的论文分析总结报告.
        第一节：Introduction
        第二节：Methods
        第三节：Results & Conclusion
        第四节：根据生成的论文分析用户可能提出的后续问题
        「写作内容」

        第一节（Introduction）

        请基于论文的引言部分（Introduction）与大模型自身的知识库信息，阐述该论文工作的研究背景与重要性。
        说明该研究如何与当前 AI 领域的发展趋势或前沿技术需求相关，并指出它所试图解决的问题。
        最终目标是让读者理解：“为什么这篇论文的研究工作重要？”
        第二节（Methods）

        请详细列举论文中使用的方法或算法流程，并对论文中给出的所有关键公式进行基于 LaTeX 的公式展示。
        在解释每个方法或公式时，务必给出溯源信息：例如可引用论文的章节、段落或语句（如“基于论文第 3.2 节第 2 段可知…”、“如作者在第 4.1 小节中所述…”等），以确保分析内容有据可依。
        
        在方法分析部分，需要解释这些方法在论文中解决了哪些具体问题或有哪些理论支撑。可以结合你自身的知识库对论文中的方法进行适当延伸或补充说明，但请与论文原文内容保持一致。
        此处的方法分析部分要详细，目标是让读者理解论文中算法的细节和实施过程，给出论文完整方法框架，完整叙述论文中使用的全部方法。梳理论文中理论方法的使用逻辑。
        第三节（Results & Conclusion）

        请先详细描述论文的结果，包括实验结果、与其他对比方法或工作之间的性能对比（如对比图表、数据结果等）。如果论文中提供了可视化图表或指标对比，请对这些内容进行简洁而准确的概括。
        在总结部分，请结合你自身的知识库，对该论文工作的意义与前景进行发散性的探讨：
        它在弥补哪些已有研究工作的不足？
        它在哪些方面具有潜在的学术价值或应用价值？
        将来在哪些领域或场景中可能会衍生出更深入或更有影响力的研究？
        强调展望内容的专业性与创新性，为读者提供对该论文研究的更深层次理解与启示。
        第四节 (后续问题)

        根据生成的论文分析，用户可能追问的后续问题。

        「写作要求」
        1. 各部分内容需详尽分析论文原文，并结合大模型知识库给出专业见解。
        2. 在分析过程中，请明确指出你的回答中哪些句子或段落与论文中的具体图片相关。例如：
        - 当在方法部分解释某算法或公式时，如果论文中提供了示意图，请在该描述后添加图片索引标注（如“图2”、“图3”），以便程序后续根据索引插入对应图片。
        - 请确保每个引用的图片都有唯一的索引，且引用信息要准确对应论文中的图例。
        3. 每个陈述句结尾必须用引用上下文，在引用上下文中的内容时，请使用 [chunk_index] 格式，例如 [3] 或 [1,5]。
        4. 全文要求保持专业性、学术风格和逻辑严谨性，关键结论或分析部分可使用编号或要点分条描述。
        
        [输入与输出形式]

        输入：我将提供给你论文的核心内容（如引言段落、研究方法、实验结果等。
        输出：基于论文内容及大模型知识库的一份专业 AI 领域论文分析总结报告，报告中必须覆盖上述所有写作要求，并在涉及图片描述的地方标注出相应的图片索引，如“图1”、“图2”等。
            请务必在回答时，针对每个与图片有关的描述都在适当位置明确标注出对应的图片索引和RAG过程的chunk索引，从而方便后续程序根据索引将提取的图片和chunk插入至报告中。
            
        """
        # 示例格式：
        #     "该研究提出了新型神经网络架构[12]，如图1所示，在ImageNet数据集上达到92.4%准确率[15]。"
        retriever = self.vectorstore.as_retriever(search_kwargs={"k": 40})
        prompt = ChatPromptTemplate.from_template(prompt_template)
        
        self.rag_chain = (
            {"context": retriever | self._format_context}
            | prompt
            | self.llm
            | StrOutputParser()
            | RunnableLambda(lambda x: print(f"LLM raw output: {x}") or x)  # 添加日志
            # | RunnableLambda(self._post_process)
        )
        return self.rag_chain.invoke("请详细分析论文HiRT: Enhancing Robotic Control with Hierarchical Robot Transformers")

    def _format_context(self, docs: List[Document]) -> str:
        """动态内容访问方法"""
        # 检测可用内容属性
        content_attr = "text" if hasattr(docs[0], "text") else "page_content"
    # 添加日志
        print(f"Retrieved docs metadata in _format_context: {[doc.metadata for doc in docs]}")
        return "\n\n".join(
            f"[Chunk {doc.metadata['chunk_index']}] {getattr(doc, content_attr)}"
            for doc in docs
        )

    # def _post_process(self, text: str) -> str:
    #     """后处理验证"""
    #     # 获取有效chunk索引范围
    #     valid_indices = set(str(i) for i in range(len(self.processed_chunks)))
    #     print(f"Valid chunk indices: {valid_indices}")
    #     print(f"Raw text before processing: {text}")
        
    #     matches = re.finditer(r"\[\d+(?:,\d+)*\]", text)
    #     matched_refs = [match.group(0) for match in matches]
    #     print(f"Matched chunk references: {matched_refs}")
        
    #     text = self._validate_refs(text, r"\[\d+(?:,\d+)*\]", valid_indices)
    #     valid_figures = set(self.figures.keys())
    #     return self._validate_refs(text, r"图\d+", valid_figures)

    # def _validate_refs(self, text: str, pattern: str, valid: set) -> str:
    #     """通用验证函数"""
    #     matches = re.finditer(pattern, text)
    #     for match in matches:
    #         ref = match.group(0)
    #         if ref not in valid:
    #             text = text.replace(ref, "[无效引用]")
    #     return text

    def save_report(self, report: str, output_path: str):
        """多格式保存"""
        ext = os.path.splitext(output_path)[1].lower()
        
        if ext == ".md":
            self._save_markdown(report, output_path)
        elif ext == ".docx":
            self._save_word(report, output_path)
        else:
            raise ValueError("仅支持.md和.docx格式")

    def _save_markdown(self, report: str, path: str):
        """Markdown保存"""
        with open(path, "w", encoding="utf-8") as f:
            f.write("# 论文分析报告\n\n")
            f.write("## 详细分析\n")
            
            # 使用正则表达式替换图片引用为图片插入代码
            figure_pattern = re.compile(r"(图\d+)")
            def replace_fig(match):
                fig_key = match.group(1)
                return self._insert_images(fig_key, is_md=True)
            
            # 替换正文中的图片引用并写入
            processed_report = figure_pattern.sub(replace_fig, report)
            f.write(processed_report)

    def _save_word(self, report: str, path: str):
        """Word保存"""
        doc = Document()
        doc.add_heading("论文分析报告", 0)
        self._insert_images(doc, report)
        doc.add_heading("详细分析", level=1)
        doc.add_paragraph(report)
        doc.save(path)

    def _insert_images(self, fig_key: str, is_md=False):
        """返回图片插入代码"""
        if fig_key in self.figures:
            if is_md:
                return f"![{fig_key}]({os.path.basename(self.figures[fig_key])})\n\n"
            else:
                # 对于 Word 文档，返回占位符
                return f"{{{{{fig_key}}}}}"
        return ""

if __name__ == "__main__":
    load_dotenv()
    
    # 配置参数
    PDF_PATH = "D:/G/paper4.pdf"
    OUTPUT_DIR = "D:/G/extracted_images"
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # 初始化分析器
    analyzer = EnhancedPaperAnalyzer(PDF_PATH)
    
    # 执行完整流程
    analyzer.process_document(OUTPUT_DIR)
    report = analyzer.generate_analysis()
    
    # 保存结果
    analyzer.save_report(report, os.path.join(OUTPUT_DIR, "analysis.md"))
    # analyzer.save_report(report, os.path.join(OUTPUT_DIR, "analysis.docx"))

    print("分析流程完成")