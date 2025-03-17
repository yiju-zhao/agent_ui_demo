import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd

def save_markdown_report(markdown_text, output_file="conference_daily_report.md"):
    """
    将Markdown文本保存到文件
    
    Args:
        markdown_text (str): Markdown格式的文本
        output_file (str): 输出文件路径
    """
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(markdown_text)
    print(f"每日参会快报已保存到 {output_file}")


def generate_daily_conference_report(df):
    """
    根据DataFrame生成每日参会快报的Markdown格式文本
    
    Args:
        df (DataFrame): 包含会议数据的DataFrame
        
    Returns:
        str: Markdown格式的每日参会快报
    """
    # 创建Markdown文本
    markdown_text = f"# 每日参会快报\n\n"
    
    # 遍历每一行生成报告内容
    for _, row in df.iterrows():
        title = row.get('标题\nTitle', '未知标题')
        session_type = row.get('Session Type', '未知会话类型')
        topic = row.get('Topic', '')
        
        # 获取实事描述和对公司启示
        facts_desc = row.get('实事描述\nDescription of Facts merged', '')
        if pd.isna(facts_desc) or not facts_desc or (isinstance(facts_desc, str) and facts_desc.isspace()):
            facts_desc = row.get('实事描述\nDescription of Facts', '')
            try:
                facts_data = eval(facts_desc) if isinstance(facts_desc, str) else facts_desc
                if isinstance(facts_data, list):
                    facts_desc = "\n".join([f"{item}" for item in facts_data])
            except:
                pass
        
        company_insights = row.get('对公司启示\nInsights for Company merged', '')
        if pd.isna(company_insights) or not company_insights or (isinstance(company_insights, str) and company_insights.isspace()):
            company_insights = row.get('对公司启示\nInsights for Company', '')
            try:
                insights_data = eval(company_insights) if isinstance(company_insights, str) else company_insights
                if isinstance(insights_data, list):
                    company_insights = "\n".join([f"{item}" for item in insights_data])
            except:
                pass
        
        composer = row.get('撰稿人\nAuthors', '')
        formatted_composer = format_composer(composer)
        
        speakers = row.get('Speakers', '')
        formatted_speakers = format_speakers(speakers)
        
        if not pd.isna(title) and title and not pd.isna(session_type) and session_type:
            markdown_text += f"## 【{session_type}】{title}\n\n"
        if not pd.isna(topic) and topic:
            markdown_text += f"### 主题\n{topic}\n\n"
        if formatted_speakers != "无":
            markdown_text += f"### 演讲人或相关公司\n{formatted_speakers}\n\n"
        if not pd.isna(facts_desc) and facts_desc:
            markdown_text += f"### 实事描述\n{facts_desc}\n\n"
        if not pd.isna(company_insights) and company_insights:
            markdown_text += f"### 对华为的启示\n{company_insights}\n\n"
        if formatted_composer != "未知撰稿人":
            markdown_text += f"撰稿人：{formatted_composer}\n\n"
        
        markdown_text += "---\n\n"
    
    return markdown_text



def format_composer(composer):
    """
    格式化撰稿人信息
    
    Args:
        composer: 原始撰稿人数据
        
    Returns:
        str: 格式化后的撰稿人字符串
    """
    formatted_composer = ""
    if composer:
        try:
            composer_data = json.loads(composer) if isinstance(composer, str) else composer
            if isinstance(composer_data, list):
                composer_items = []
                for person in composer_data:
                    if isinstance(person, dict):
                        name = person.get('name', '')
                        id_num = person.get('id', '')
                        if name and id_num:
                            composer_items.append(f"{name} {id_num}")
                formatted_composer = "、".join(composer_items)
            elif isinstance(composer_data, dict):
                name = composer_data.get('name', '')
                id_num = composer_data.get('id', '')
                if name and id_num:
                    formatted_composer = f"{name} {id_num}"
        except:
            formatted_composer = composer
    
    return formatted_composer if formatted_composer else "未知撰稿人"


def format_speakers(speakers):
    """
    格式化演讲者信息
    
    Args:
        speakers: 原始演讲者数据
        
    Returns:
        str: 格式化后的演讲者字符串
    """
    formatted_speakers = "无"
    if speakers:
        try:
            speaker_data = json.loads(speakers) if isinstance(speakers, str) else speakers
            if isinstance(speaker_data, list):
                speaker_items = []
                for person in speaker_data:
                    if isinstance(person, dict):
                        name = person.get('name', '')
                        position = person.get('position', '')
                        company = person.get('company', '')
                        
                        other_info = []
                        if position:
                            other_info.append(position)
                        if company:
                            other_info.append(company)
                        
                        if name:
                            if other_info:
                                speaker_items.append(f"{name}（{', '.join(other_info)}）")
                            else:
                                speaker_items.append(name)
                
                if speaker_items:
                    formatted_speakers = "、".join(speaker_items)
            elif isinstance(speaker_data, dict):
                name = speaker_data.get('name', '')
                position = speaker_data.get('position', '')
                company = speaker_data.get('company', '')
                
                other_info = []
                if position:
                    other_info.append(position)
                if company:
                    other_info.append(company)
                
                if name:
                    if other_info:
                        formatted_speakers = f"{name}（{', '.join(other_info)}）"
                    else:
                        formatted_speakers = name
        except:
            formatted_speakers = speakers
    
    return formatted_speakers


def markdown_to_word(markdown_text, output_file="conference_daily_report_from_md.docx"):
    """
    将Markdown文本直接转换为Word文档
    
    Args:
        markdown_text (str): Markdown格式的文本
        output_file (str): 输出的Word文件路径
    """
    import markdown
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 创建Word文档
    doc = Document()
    
    # 将Markdown转换为HTML
    html = markdown.markdown(markdown_text)
    
    # 使用BeautifulSoup解析HTML
    soup = BeautifulSoup(html, 'html.parser')
    
    # 处理标题
    for h1 in soup.find_all('h1'):
        heading = doc.add_heading(h1.text, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 处理段落和其他元素
    current_element = soup.find('h1')
    if current_element:
        current_element = current_element.find_next()
    else:
        current_element = soup.find_next()
    
    while current_element:
        if current_element.name == 'h2':
            # 二级标题
            doc.add_heading(current_element.text, level=1)
        elif current_element.name == 'h3':
            # 三级标题
            doc.add_heading(current_element.text, level=2)
        elif current_element.name == 'p':
            # 普通段落
            if current_element.text.strip() == '---':
                # 处理分隔线
                doc.add_paragraph('_' * 40)
            else:
                # 处理普通文本
                doc.add_paragraph(current_element.text)
        elif current_element.name == 'ul':
            # 处理无序列表
            for li in current_element.find_all('li'):
                p = doc.add_paragraph()
                p.add_run('• ' + li.text)
        
        # 移动到下一个元素
        current_element = current_element.find_next()
    
    # 保存文档
    doc.save(output_file)
    print(f"从Markdown生成的Word文档已保存到 {output_file}") 